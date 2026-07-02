from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest import mock
from io import StringIO

from docx import Document

import finalize_subs
import rename_subs


class RenameSubsTest(unittest.TestCase):
    def test_final_name_replaces_al_el_suffix(self) -> None:
        path = Path("人文講堂_其實你胖得很冤枉  - 陳欣湄 [2]_al_el.docx")

        self.assertEqual(
            rename_subs.final_name_for(path),
            Path("人文講堂_其實你胖得很冤枉  - 陳欣湄 [2]_final.docx"),
        )

    def test_final_name_replaces_al_sy_suffix(self) -> None:
        path = Path("sample_al_sy.docx")

        self.assertEqual(rename_subs.final_name_for(path), Path("sample_final.docx"))

    def test_rename_docx_moves_file_to_final_suffix(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample_al_el.docx"
            source.touch()

            destination = rename_subs.rename_docx(source)

            self.assertEqual(destination, Path(tmp_dir) / "sample_final.docx")
            self.assertFalse(source.exists())
            self.assertTrue(destination.exists())


class FinalizeSubsTest(unittest.TestCase):
    def test_finalize_docx_cleans_then_renames(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample_al_el.docx"
            Document().save(source)

            def fake_clean(input_path: Path, output_path: Path) -> None:
                self.assertEqual(input_path, source)
                self.assertEqual(output_path, source)
                doc = Document()
                doc.add_paragraph("字幕：")
                doc.add_paragraph("00:00:01:00\t00:00:02:00\t中文")
                doc.add_paragraph("English line.")
                doc.save(output_path)

            with (
                mock.patch("finalize_subs.clean_subs.remove_sources_from_docx", side_effect=fake_clean),
                mock.patch("finalize_subs.thumbnail_subs.export_thumbnail_from_docx") as export_thumbnail,
                mock.patch("finalize_subs.rename_subs.rename_docx") as rename_docx,
            ):
                export_thumbnail.return_value = Path(tmp_dir) / "sample.png"
                rename_docx.side_effect = lambda path: Path(tmp_dir) / "sample_final.docx"

                result = finalize_subs.finalize_docx(source)

            self.assertEqual(result.final_path, Path(tmp_dir) / "sample_final.docx")
            self.assertEqual(result.thumbnail_path, Path(tmp_dir) / "sample.png")
            self.assertEqual(
                result.analysis_text,
                "00:00:01:00\t00:00:02:00\t中文\nEnglish line.\n",
            )
            export_thumbnail.assert_called_once_with(source)
            rename_docx.assert_called_once_with(source)

    def test_extract_subtitle_analysis_text_reads_only_subtitle_section(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            doc = Document()
            doc.add_paragraph("建議YT標題：")
            doc.add_paragraph("Title")
            doc.add_paragraph("字幕：")
            doc.add_paragraph("00:00:01:00\t00:00:02:00\t中文")
            doc.add_paragraph("English line.")
            doc.add_paragraph("")
            doc.add_paragraph("00:00:02:00\t00:00:03:00\t中文二")
            doc.add_paragraph("Second line.")
            doc.add_paragraph("簡介：")
            doc.add_paragraph("Not analyzed")
            doc.save(source)

            self.assertEqual(
                finalize_subs.extract_subtitle_analysis_text(source),
                "\n".join(
                    [
                        "00:00:01:00\t00:00:02:00\t中文",
                        "English line.",
                        "00:00:02:00\t00:00:03:00\t中文二",
                        "Second line.",
                        "",
                    ]
                ),
            )

    def test_main_reports_renamed_not_finalized(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample_al_el.docx"
            source.touch()
            stdout = StringIO()

            with (
                mock.patch("sys.stdout", stdout),
                mock.patch(
                    "finalize_subs.finalize_docx",
                    return_value=finalize_subs.FinalizeResult(
                        final_path=Path(tmp_dir) / "sample_final.docx",
                        thumbnail_path=Path(tmp_dir) / "sample.png",
                        analysis_text="00:00:01:00\t00:00:02:00\t中文\nEnglish line.\n",
                    ),
                ),
                mock.patch("finalize_subs.run_subtitle_analysis", return_value=0) as run_analysis,
            ):
                exit_code = finalize_subs.main([str(source)])

            self.assertEqual(exit_code, 0)
            run_analysis.assert_called_once_with(
                [
                    (
                        Path(tmp_dir) / "sample_final.docx",
                        "00:00:01:00\t00:00:02:00\t中文\nEnglish line.\n",
                    )
                ]
            )
            output = stdout.getvalue()
            self.assertIn("[cleaned]", output)
            self.assertIn("[exported]", output)
            self.assertIn("[renamed]", output)
            self.assertNotIn("[finalized]", output)

    def test_run_subtitle_analysis_invokes_watch_once(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            watch_ts = tmp_path / "node" / "sub" / "src" / "cli" / "watch.ts"
            watch_ts.parent.mkdir(parents=True)
            watch_ts.touch()
            source_path = tmp_path / "sample_final.docx"

            completed = mock.Mock(returncode=0)
            captured_text = []

            def fake_run(argv, cwd):
                captured_text.append(Path(argv[4]).read_text(encoding="utf-8"))
                return completed

            with (
                mock.patch.dict("os.environ", {"SUB_WATCH_TS": str(watch_ts)}),
                mock.patch("finalize_subs.subprocess.run", side_effect=fake_run) as run,
            ):
                exit_code = finalize_subs.run_subtitle_analysis(
                    [(source_path, "00:00:01:00\t00:00:02:00\t中文\nEnglish line.\n")]
                )

            self.assertEqual(exit_code, 0)
            argv = run.call_args.args[0]
            self.assertEqual(argv[:4], ["npx", "tsx", str(watch_ts), "--once"])
            self.assertEqual(argv[5:], ["--type", "subs"])
            self.assertEqual(captured_text, ["00:00:01:00\t00:00:02:00\t中文\nEnglish line.\n"])
            self.assertEqual(run.call_args.kwargs["cwd"], tmp_path / "node" / "sub")

    def test_wrapper_and_symlink_split_matches_dependency_needs(self) -> None:
        repo_root = Path(__file__).resolve().parents[1]

        rename_command = repo_root / "rename-subs"
        finalize_wrapper = (repo_root / "finalize-subs").read_text(encoding="utf-8")

        self.assertTrue(rename_command.is_symlink())
        self.assertEqual(rename_command.resolve(), (repo_root / "rename_subs.py").resolve())
        self.assertIn('"$HOME/python/word/.venv/bin/python"', finalize_wrapper)
        self.assertIn('"$HOME/python/subs-tools/finalize_subs.py"', finalize_wrapper)


if __name__ == "__main__":
    unittest.main()
