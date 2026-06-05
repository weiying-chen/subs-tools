import tempfile
import unittest
from pathlib import Path
import xml.etree.ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

import clean_subs


class CleanSubsHighlightTest(unittest.TestCase):
    def test_strips_run_shading_from_document_xml_bytes(self) -> None:
        xml = b"""<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:r>
        <w:rPr>
          <w:shd w:val="clear" w:color="auto" w:fill="00B0F0"/>
          <w:highlight w:val="yellow"/>
        </w:rPr>
        <w:t>text</w:t>
      </w:r>
    </w:p>
  </w:body>
</w:document>
"""
        cleaned = clean_subs._strip_run_shading_xml(xml)
        root = ET.fromstring(cleaned)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        self.assertEqual(len(root.findall(".//w:shd", ns)), 0)
        self.assertEqual(len(root.findall(".//w:highlight", ns)), 1)

    def _write_docx_with_highlights(
        self,
        path: Path,
        paragraphs: list[list[tuple[str, WD_COLOR_INDEX | None]]],
    ) -> None:
        doc = Document()
        for run_specs in paragraphs:
            paragraph = doc.add_paragraph()
            for text, color in run_specs:
                run = paragraph.add_run(text)
                run.font.highlight_color = color
        doc.save(path)

    def _set_run_shading(self, path: Path, paragraph_index: int, run_index: int, fill: str) -> None:
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        with ZipFile(path, "r") as zin:
            document_xml = zin.read("word/document.xml")
            entries = [(info, zin.read(info.filename)) for info in zin.infolist()]
        tree = ET.ElementTree(ET.fromstring(document_xml))
        root = tree.getroot()
        paragraphs = root.findall(".//w:body/w:p", ns)
        run = paragraphs[paragraph_index].findall("w:r", ns)[run_index]
        rpr = run.find("w:rPr", ns)
        if rpr is None:
            rpr = ET.SubElement(run, "{%s}rPr" % ns["w"])
        shd = ET.SubElement(rpr, "{%s}shd" % ns["w"])
        shd.set("{%s}val" % ns["w"], "clear")
        shd.set("{%s}color" % ns["w"], "auto")
        shd.set("{%s}fill" % ns["w"], fill)
        updated_document_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        with ZipFile(path, "w", compression=ZIP_DEFLATED) as zout:
            for info, data in entries:
                if info.filename == "word/document.xml":
                    data = updated_document_xml
                zout.writestr(info, data)

    def test_removes_all_highlights_except_restored_yellow_on_parenthesis_lines(self) -> None:
        temp_dir = Path(tempfile.mkdtemp(prefix="subs_tools_clean_test_"))
        source_path = temp_dir / "input.docx"
        output_path = temp_dir / "output.docx"
        self._write_docx_with_highlights(
            source_path,
            [
                [
                    ("Regular highlighted line", WD_COLOR_INDEX.BRIGHT_GREEN),
                ],
                [
                    ("(", WD_COLOR_INDEX.YELLOW),
                    ("Parenthesized note", None),
                    (")", WD_COLOR_INDEX.TURQUOISE),
                ],
                [
                    ("(", None),
                    ("No original yellow here", WD_COLOR_INDEX.TURQUOISE),
                    (")", None),
                ],
            ],
        )

        clean_subs.remove_sources_from_docx(source_path, output_path)

        out_doc = Document(output_path)
        first_colors = [run.font.highlight_color for run in out_doc.paragraphs[0].runs]
        second_colors = [run.font.highlight_color for run in out_doc.paragraphs[1].runs]
        third_colors = [run.font.highlight_color for run in out_doc.paragraphs[2].runs]

        self.assertEqual(first_colors, [None])
        self.assertEqual(
            second_colors,
            [WD_COLOR_INDEX.YELLOW, WD_COLOR_INDEX.YELLOW, WD_COLOR_INDEX.YELLOW],
        )
        self.assertEqual(third_colors, [None, None, None])

    def test_preserves_yellow_for_split_parenthesis_lines(self) -> None:
        temp_dir = Path(tempfile.mkdtemp(prefix="subs_tools_clean_test_"))
        source_path = temp_dir / "input.docx"
        output_path = temp_dir / "output.docx"
        self._write_docx_with_highlights(
            source_path,
            [
                [
                    ("(Dr. Chuang Chia-ying,", WD_COLOR_INDEX.YELLOW),
                ],
                [
                    ("Taichung Tzu Chi Hospital", None),
                    (")", WD_COLOR_INDEX.YELLOW),
                ],
            ],
        )

        clean_subs.remove_sources_from_docx(source_path, output_path)

        out_doc = Document(output_path)
        first_colors = [run.font.highlight_color for run in out_doc.paragraphs[0].runs]
        second_colors = [run.font.highlight_color for run in out_doc.paragraphs[1].runs]

        self.assertEqual(first_colors, [WD_COLOR_INDEX.YELLOW])
        self.assertEqual(second_colors, [WD_COLOR_INDEX.YELLOW, WD_COLOR_INDEX.YELLOW])

    def test_preserves_yellow_on_timestamp_paragraph_before_parenthesis_note(self) -> None:
        temp_dir = Path(tempfile.mkdtemp(prefix="subs_tools_clean_test_"))
        source_path = temp_dir / "input.docx"
        output_path = temp_dir / "output.docx"
        self._write_docx_with_highlights(
            source_path,
            [
                [
                    ("00:02:10:25", WD_COLOR_INDEX.YELLOW),
                    ("\t", None),
                    ("00:02:13:26", WD_COLOR_INDEX.YELLOW),
                    ("\t", None),
                    ("邪積胸中", WD_COLOR_INDEX.YELLOW),
                    (" ", None),
                    ("阻塞氣逆", WD_COLOR_INDEX.YELLOW),
                ],
                [
                    ("(", WD_COLOR_INDEX.YELLOW),
                    ("Pathogenic Buildup in the Chest", None),
                    (")", WD_COLOR_INDEX.YELLOW),
                ],
            ],
        )

        clean_subs.remove_sources_from_docx(source_path, output_path)

        out_doc = Document(output_path)
        first_colors = [run.font.highlight_color for run in out_doc.paragraphs[0].runs]
        second_colors = [run.font.highlight_color for run in out_doc.paragraphs[1].runs]

        self.assertEqual(
            first_colors,
            [
                WD_COLOR_INDEX.YELLOW,
                WD_COLOR_INDEX.YELLOW,
                WD_COLOR_INDEX.YELLOW,
                WD_COLOR_INDEX.YELLOW,
                WD_COLOR_INDEX.YELLOW,
                WD_COLOR_INDEX.YELLOW,
                WD_COLOR_INDEX.YELLOW,
            ],
        )
        self.assertEqual(
            second_colors,
            [WD_COLOR_INDEX.YELLOW, WD_COLOR_INDEX.YELLOW, WD_COLOR_INDEX.YELLOW],
        )

    def test_removes_run_shading_colors_from_non_parenthesis_lines(self) -> None:
        temp_dir = Path(tempfile.mkdtemp(prefix="subs_tools_clean_test_"))
        source_path = temp_dir / "input.docx"
        output_path = temp_dir / "output.docx"
        self._write_docx_with_highlights(
            source_path,
            [
                [("which is a non-small cell cancer.", None)],
                [("certain ", None), ("cells", None), (" near the lungs turn cancerous.", None)],
            ],
        )
        self._set_run_shading(source_path, 0, 0, "00B0F0")
        self._set_run_shading(source_path, 1, 0, "92D050")
        self._set_run_shading(source_path, 1, 1, "92D050")

        clean_subs.remove_sources_from_docx(source_path, output_path)

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        with ZipFile(output_path, "r") as zin:
            root = ET.fromstring(zin.read("word/document.xml"))
        self.assertEqual(len(root.findall(".//w:shd", ns)), 0)


if __name__ == "__main__":
    unittest.main()
