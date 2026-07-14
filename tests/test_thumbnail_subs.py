from __future__ import annotations

import base64
import tempfile
import unittest
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
import xml.etree.ElementTree as ET

from docx import Document

import thumbnail_subs


PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAFgwJ/"
    "l6NqSAAAAABJRU5ErkJggg=="
)


class ThumbnailSubsTest(unittest.TestCase):
    def _write_docx_with_title_and_image(self, path: Path) -> None:
        image_path = path.with_suffix(".png")
        image_path.write_bytes(PNG_BYTES)
        doc = Document()
        doc.add_paragraph("建議YT標題：")
        doc.add_paragraph(
            "Chinese Medicine Clinic - How TCM Helps Prevent Cancer "
            "(大愛學漢醫 - 陽虛、血瘀、痰濕體質 防癌作法)"
        )
        doc.add_picture(str(image_path))
        doc.save(path)

    def test_thumbnail_stem_uses_english_part_of_youtube_title(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "sample.docx"
            self._write_docx_with_title_and_image(docx_path)

            self.assertEqual(
                thumbnail_subs.thumbnail_stem_for_docx(docx_path),
                "Chinese Medicine Clinic - How TCM Helps Prevent Cancer",
            )

    def test_thumbnail_stem_omits_invalid_title_punctuation(self) -> None:
        self.assertEqual(
            thumbnail_subs.sanitize_filename("Does Freedom Shape Our Sense of Beauty?"),
            "Does Freedom Shape Our Sense of Beauty",
        )

    def test_export_thumbnail_uses_first_referenced_document_image(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "sample.docx"
            self._write_docx_with_title_and_image(docx_path)

            output_path = thumbnail_subs.export_thumbnail_from_docx(docx_path)

            self.assertEqual(
                output_path,
                Path(tmp_dir) / "Chinese Medicine Clinic - How TCM Helps Prevent Cancer.png",
            )
            self.assertEqual(output_path.read_bytes(), PNG_BYTES)

    def test_export_thumbnail_ignores_unreferenced_media_files(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "sample.docx"
            self._write_docx_with_title_and_image(docx_path)
            stale_bytes = b"stale image"
            with ZipFile(docx_path, "a", compression=ZIP_DEFLATED) as docx:
                docx.writestr("word/media/stale.png", stale_bytes)

            output_path = thumbnail_subs.export_thumbnail_from_docx(docx_path)

            self.assertNotEqual(output_path.read_bytes(), stale_bytes)

    def test_export_thumbnail_supports_vml_image_references(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "sample.docx"
            self._write_docx_with_title_and_image(docx_path)
            ns = {
                "a": thumbnail_subs.DRAWING_NS,
                "r": thumbnail_subs.REL_NS,
                "w": thumbnail_subs.WORD_NS,
                "v": thumbnail_subs.VML_NS,
            }
            with ZipFile(docx_path, "r") as docx:
                entries = [(info, docx.read(info.filename)) for info in docx.infolist()]
                document_xml = docx.read("word/document.xml")
            root = ET.fromstring(document_xml)
            blip = root.find(".//a:blip", ns)
            rel_id = blip.get(f"{{{thumbnail_subs.REL_NS}}}embed")
            body = root.find(".//w:body", ns)
            for paragraph in list(body.findall("w:p", ns)):
                if paragraph.find(".//a:blip", ns) is not None:
                    body.remove(paragraph)
            paragraph = ET.SubElement(body, f"{{{thumbnail_subs.WORD_NS}}}p")
            run = ET.SubElement(paragraph, f"{{{thumbnail_subs.WORD_NS}}}r")
            pict = ET.SubElement(run, f"{{{thumbnail_subs.WORD_NS}}}pict")
            image_data = ET.SubElement(pict, f"{{{thumbnail_subs.VML_NS}}}imagedata")
            image_data.set(f"{{{thumbnail_subs.REL_NS}}}id", rel_id)
            updated_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            with ZipFile(docx_path, "w", compression=ZIP_DEFLATED) as docx:
                for info, data in entries:
                    if info.filename == "word/document.xml":
                        data = updated_xml
                    docx.writestr(info, data)

            output_path = thumbnail_subs.export_thumbnail_from_docx(docx_path)

            self.assertEqual(output_path.read_bytes(), PNG_BYTES)


if __name__ == "__main__":
    unittest.main()
