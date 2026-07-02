from io import BytesIO
from pathlib import Path
import tempfile
import unittest
from unittest import mock
from zipfile import ZIP_DEFLATED, ZipFile
import xml.etree.ElementTree as ET

import clean_subs
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _highlight_run(run, value: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), value)
    r_pr.append(highlight)


class CleanSubsCliTest(unittest.TestCase):
    def test_clean_subs_wrapper_uses_home_word_venv(self):
        wrapper_path = Path(__file__).resolve().parents[1] / "clean-subs"
        wrapper = wrapper_path.read_text(encoding="utf-8")

        self.assertIn('"$HOME/python/word/.venv/bin/python"', wrapper)
        self.assertIn('"$HOME/python/subs-tools/clean_subs.py"', wrapper)
        self.assertNotIn('"$ROOT/.venv/bin/python"', wrapper)

    def test_resolve_input_paths_uses_current_directory_docx_when_empty(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            first = tmp_path / "a.docx"
            second = tmp_path / "b.docx"
            ignored = tmp_path / "notes.txt"
            first.touch()
            second.touch()
            ignored.touch()

            with mock.patch("clean_subs.Path.cwd", return_value=tmp_path):
                self.assertEqual(
                    clean_subs._resolve_input_paths([]),
                    [first, second],
                )

    def test_resolve_input_paths_keeps_explicit_files(self):
        self.assertEqual(
            clean_subs._resolve_input_paths(["one.docx", "two.docx"]),
            [Path("one.docx"), Path("two.docx")],
        )

    def test_remove_sources_preserves_shaded_paragraphs(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            input_path = Path(tmp_dir) / "input.docx"
            output_path = Path(tmp_dir) / "output.docx"
            doc = Document()
            doc.add_paragraph("keep this line")
            shaded = doc.add_paragraph("keep this marked line")
            _shade_paragraph(shaded, "FF0000")
            doc.add_paragraph("keep this too")
            doc.save(input_path)

            clean_subs.remove_sources_from_docx(input_path, output_path)

            cleaned_text = "\n".join(
                paragraph.text for paragraph in Document(output_path).paragraphs
            )
            self.assertIn("keep this line", cleaned_text)
            self.assertIn("keep this marked line", cleaned_text)
            self.assertIn("keep this too", cleaned_text)

    def test_remove_sources_strips_highlight_and_shading_formatting(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            input_path = Path(tmp_dir) / "input.docx"
            output_path = Path(tmp_dir) / "output.docx"
            doc = Document()
            summary = doc.add_paragraph()
            summary_run = summary.add_run("(1)04:10-06:52 (2m42s)")
            _highlight_run(summary_run, "yellow")
            doc.add_paragraph("字幕：")
            shaded = doc.add_paragraph("keep shaded text")
            _shade_paragraph(shaded, "FF0000")
            highlighted = doc.add_paragraph()
            run = highlighted.add_run("keep highlighted text")
            _highlight_run(run, "yellow")
            doc.save(input_path)

            clean_subs.remove_sources_from_docx(input_path, output_path)

            output_doc = Document(output_path)
            cleaned_text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
            self.assertIn("(1)04:10-06:52 (2m42s)", cleaned_text)
            self.assertIn("keep shaded text", cleaned_text)
            self.assertIn("keep highlighted text", cleaned_text)
            self.assertTrue(output_doc.paragraphs[0]._p.findall(".//w:highlight", {"w": clean_subs.WORD_NAMESPACE}))
            self.assertFalse(output_doc.paragraphs[2]._p.findall(".//w:shd", {"w": clean_subs.WORD_NAMESPACE}))
            self.assertFalse(output_doc.paragraphs[2]._p.findall(".//w:highlight", {"w": clean_subs.WORD_NAMESPACE}))
            self.assertFalse(output_doc.paragraphs[3]._p.findall(".//w:shd", {"w": clean_subs.WORD_NAMESPACE}))
            self.assertTrue(output_doc.paragraphs[3]._p.findall(".//w:highlight", {"w": clean_subs.WORD_NAMESPACE}))

    def test_remove_sources_strips_intro_marking_formatting(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            input_path = Path(tmp_dir) / "input.docx"
            output_path = Path(tmp_dir) / "output.docx"
            doc = Document()
            summary = doc.add_paragraph()
            summary_run = summary.add_run("(1)04:10-06:52 (2m42s)")
            _highlight_run(summary_run, "yellow")
            doc.add_paragraph("簡介：")
            shaded = doc.add_paragraph("keep intro shaded text")
            _shade_paragraph(shaded, "FF0000")
            highlighted = doc.add_paragraph()
            run = highlighted.add_run("keep intro highlighted text")
            _highlight_run(run, "yellow")
            doc.save(input_path)

            clean_subs.remove_sources_from_docx(input_path, output_path)

            output_doc = Document(output_path)
            cleaned_text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
            self.assertIn("(1)04:10-06:52 (2m42s)", cleaned_text)
            self.assertIn("keep intro shaded text", cleaned_text)
            self.assertIn("keep intro highlighted text", cleaned_text)
            self.assertTrue(output_doc.paragraphs[0]._p.findall(".//w:highlight", {"w": clean_subs.WORD_NAMESPACE}))
            self.assertFalse(output_doc.paragraphs[2]._p.findall(".//w:shd", {"w": clean_subs.WORD_NAMESPACE}))
            self.assertFalse(output_doc.paragraphs[2]._p.findall(".//w:highlight", {"w": clean_subs.WORD_NAMESPACE}))
            self.assertFalse(output_doc.paragraphs[3]._p.findall(".//w:shd", {"w": clean_subs.WORD_NAMESPACE}))
            self.assertTrue(output_doc.paragraphs[3]._p.findall(".//w:highlight", {"w": clean_subs.WORD_NAMESPACE}))

    def test_remove_sources_strips_chatgpt_image_credit_and_collapses_blanks(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            input_path = Path(tmp_dir) / "input.docx"
            output_path = Path(tmp_dir) / "output.docx"
            doc = Document()
            doc.add_paragraph("Before")
            doc.add_paragraph("")
            doc.add_paragraph("Image created with ChatGPT.")
            doc.add_paragraph("")
            doc.add_paragraph("")
            doc.add_paragraph("After")
            doc.save(input_path)

            clean_subs.remove_sources_from_docx(input_path, output_path)

            texts = [paragraph.text for paragraph in Document(output_path).paragraphs]
            self.assertNotIn("Image created with ChatGPT.", texts)
            self.assertEqual(texts, ["Before", "", "After"])

    def test_repair_missing_use_local_dpi_namespace(self):
        broken_xml = (
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            b"<w:body><w:p><w:r><w:drawing><ns6:useLocalDpi val=\"0\"/>"
            b"</w:drawing></w:r></w:p></w:body></w:document>"
        )

        repaired_xml = clean_subs._repair_word_xml_namespaces(broken_xml)

        self.assertIn(b'xmlns:ns6="http://schemas.microsoft.com/office/drawing/2010/main"', repaired_xml)

    def test_normalize_preserves_repaired_use_local_dpi_namespace(self):
        original_xml = (
            b'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            b"<w:body/></w:document>"
        )
        cleaned_xml = (
            b'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            b'xmlns:ns6="http://schemas.microsoft.com/office/drawing/2010/main">'
            b"<w:body><ns6:useLocalDpi val=\"0\"/></w:body></w:document>"
        )

        normalized_xml = clean_subs._normalize_xml_part_against_original(
            cleaned_xml,
            original_xml,
        )

        self.assertIn(b'xmlns:ns6="http://schemas.microsoft.com/office/drawing/2010/main"', normalized_xml)

    def test_normalize_document_xml_with_declaration_keeps_valid_root(self):
        original_xml = (
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            b'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
            b'mc:Ignorable="w14 w15">'
            b"<w:body/></w:document>"
        )
        cleaned_xml = (
            b"<?xml version='1.0' encoding='utf-8'?>\n"
            b'<ns0:document xmlns:ns0="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            b'xmlns:ns2="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
            b'xmlns:ns3="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
            b'<ns0:body><ns0:p><ns0:hyperlink ns2:id="rId9"/></ns0:p></ns0:body>'
            b'<ns3:inline/>'
            b"</ns0:document>"
        )

        normalized_xml = clean_subs._normalize_xml_part_against_original(
            cleaned_xml,
            original_xml,
        )

        self.assertNotIn(b'">dy"', normalized_xml)
        self.assertNotIn(b">org/", normalized_xml)
        self.assertRegex(normalized_xml, rb">\s*<w:body>")
        self.assertNotIn(b"ns2:id", normalized_xml)
        self.assertIn(b'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"', normalized_xml)
        self.assertIn(b"r:id", normalized_xml)
        self.assertIn(b'xmlns:ns3="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"', normalized_xml)
        ET.fromstring(normalized_xml)

    def test_remove_sources_disables_track_revisions_in_settings_xml(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            input_path = Path(tmp_dir) / "input.docx"
            output_path = Path(tmp_dir) / "output.docx"
            doc = Document()
            doc.add_paragraph("keep this line")
            doc.save(input_path)

            with ZipFile(input_path, "r") as docx:
                entries = [(info, docx.read(info.filename)) for info in docx.infolist()]
            with ZipFile(input_path, "w", compression=ZIP_DEFLATED) as docx:
                for info, data in entries:
                    if info.filename == "word/settings.xml":
                        data = data.replace(
                            b"<w:defaultTabStop",
                            b"<w:trackRevisions/><w:defaultTabStop",
                            1,
                        )
                    docx.writestr(info, data)

            clean_subs.remove_sources_from_docx(input_path, output_path)

            with ZipFile(output_path, "r") as docx:
                settings_xml = docx.read("word/settings.xml")
            root = ET.fromstring(settings_xml)
            ns = {"w": clean_subs.WORD_NAMESPACE}
            self.assertEqual(root.findall(".//w:trackRevisions", ns), [])
            self.assertTrue(root.findall(".//w:defaultTabStop", ns))

    def test_remove_sources_accepts_inserted_and_deleted_text_in_output(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            input_path = Path(tmp_dir) / "input.docx"
            output_path = Path(tmp_dir) / "output.docx"
            doc = Document()
            doc.add_paragraph("placeholder")
            doc.save(input_path)

            ns = {"w": clean_subs.WORD_NAMESPACE}
            with ZipFile(input_path, "r") as zin:
                entries = [(info, zin.read(info.filename)) for info in zin.infolist()]
                document_xml = zin.read("word/document.xml")

            root = ET.fromstring(document_xml)
            paragraph = root.find(".//w:body/w:p", ns)
            for child in list(paragraph):
                paragraph.remove(child)
            inserted = ET.SubElement(paragraph, f"{{{clean_subs.WORD_NAMESPACE}}}ins")
            inserted_run = ET.SubElement(inserted, f"{{{clean_subs.WORD_NAMESPACE}}}r")
            inserted_text = ET.SubElement(inserted_run, f"{{{clean_subs.WORD_NAMESPACE}}}t")
            inserted_text.text = "accepted"
            deleted = ET.SubElement(paragraph, f"{{{clean_subs.WORD_NAMESPACE}}}del")
            deleted_run = ET.SubElement(deleted, f"{{{clean_subs.WORD_NAMESPACE}}}r")
            deleted_text = ET.SubElement(deleted_run, f"{{{clean_subs.WORD_NAMESPACE}}}delText")
            deleted_text.text = "deleted"

            updated_document_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            with ZipFile(input_path, "w", compression=ZIP_DEFLATED) as zout:
                for info, data in entries:
                    if info.filename == "word/document.xml":
                        data = updated_document_xml
                    zout.writestr(info, data)

            clean_subs.remove_sources_from_docx(input_path, output_path)

            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)
            self.assertIn("accepted", text)
            self.assertNotIn("deleted", text)
            with ZipFile(output_path, "r") as docx:
                document_xml = docx.read("word/document.xml")
            self.assertNotIn(b"<w:ins", document_xml)
            self.assertNotIn(b"<w:del", document_xml)

    def test_accepting_revisions_drops_leading_inserted_break_runs(self):
        root = ET.fromstring(
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            b"<w:body><w:p><w:pPr/><w:r><w:cr/></w:r><w:r><w:t>Summary text</w:t></w:r></w:p></w:body></w:document>"
        )

        clean_subs._strip_leading_break_runs(root)

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraph = root.find(".//w:body/w:p", ns)
        self.assertEqual(len(paragraph.findall("w:r", ns)), 1)
        self.assertEqual("".join(t.text or "" for t in paragraph.findall(".//w:t", ns)), "Summary text")


if __name__ == "__main__":
    unittest.main()
