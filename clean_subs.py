#!/usr/bin/env python3

from __future__ import annotations

import argparse
from io import BytesIO
from pathlib import Path
import re
from zipfile import ZIP_DEFLATED, ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn


INDENT_THRESHOLD_PT = 21.0
TIMESTAMP_LINE_RE = re.compile(
    r"^(?:[^\t]+\t)?\d{2}:\d{2}:\d{2}:\d{2}\t\d{2}:\d{2}:\d{2}:\d{2}\t"
)
SECTION_LABELS = {
    "建議YT標題：",
    "建議YT標題:",
    "建議標題：",
    "建議標題:",
    "簡介：",
    "簡介:",
    "選圖：",
    "選圖:",
    "字幕：",
    "字幕:",
}
SUBTITLE_LABELS = {"字幕：", "字幕:"}
YELLOW_MARKER_TEXTS = {"XXX"}
WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
OFFICE_DRAWING_NAMESPACE = "http://schemas.microsoft.com/office/drawing/2010/main"
INSERTION_TAGS = {f"{{{WORD_NAMESPACE}}}ins", f"{{{WORD_NAMESPACE}}}moveTo"}
DELETION_TAGS = {f"{{{WORD_NAMESPACE}}}del", f"{{{WORD_NAMESPACE}}}moveFrom"}
CHANGE_TAGS = {
    f"{{{WORD_NAMESPACE}}}pPrChange",
    f"{{{WORD_NAMESPACE}}}rPrChange",
    f"{{{WORD_NAMESPACE}}}tblPrChange",
    f"{{{WORD_NAMESPACE}}}tcPrChange",
    f"{{{WORD_NAMESPACE}}}trPrChange",
    f"{{{WORD_NAMESPACE}}}moveFromRangeStart",
    f"{{{WORD_NAMESPACE}}}moveFromRangeEnd",
    f"{{{WORD_NAMESPACE}}}moveToRangeStart",
    f"{{{WORD_NAMESPACE}}}moveToRangeEnd",
}
TRACK_REVISIONS_TAG = f"{{{WORD_NAMESPACE}}}trackRevisions"
REVISION_XML_BASENAMES = {
    "document.xml",
}
XMLNS_DECL_RE = re.compile(r'xmlns:([A-Za-z0-9_]+)="([^"]+)"')
TRACK_REVISIONS_XML_RE = re.compile(
    r"<(?P<prefix>[A-Za-z0-9_]+):trackRevisions\b[^>]*/>"
    r"|<(?P<prefix2>[A-Za-z0-9_]+):trackRevisions\b[^>]*>\s*</(?P=prefix2):trackRevisions>"
)
DRAWING_PREFIX_BY_URI = {
    "http://schemas.openxmlformats.org/drawingml/2006/main": "a",
    "http://schemas.openxmlformats.org/drawingml/2006/picture": "pic",
}
DOCUMENT_PREFIX_BY_URI = {
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships": "r",
    **DRAWING_PREFIX_BY_URI,
}


def _should_accept_revisions_in_part(filename: str) -> bool:
    if not filename.startswith("word/") or not filename.endswith(".xml"):
        return False
    basename = Path(filename).name
    if basename in REVISION_XML_BASENAMES:
        return True
    return basename.startswith("header") or basename.startswith("footer")


def _should_clean_settings_part(filename: str) -> bool:
    return filename == "word/settings.xml"


def _repair_word_xml_namespaces(xml_data: bytes) -> bytes:
    if b"ns6:useLocalDpi" not in xml_data or b"xmlns:ns6=" in xml_data:
        return xml_data
    return xml_data.replace(
        b"<w:document ",
        f'<w:document xmlns:ns6="{OFFICE_DRAWING_NAMESPACE}" '.encode("utf-8"),
        1,
    )


def _find_parent(root, target):
    for parent in root.iter():
        if target in list(parent):
            return parent
    return None


def _first_root_tag(xml_text: str) -> str:
    match = re.search(r"<[^?][^>]*>", xml_text)
    if match is None:
        raise ValueError("No root tag found.")
    return match.group(0)


def _namespace_prefixes(root_tag: str) -> dict[str, str]:
    return {uri: prefix for prefix, uri in XMLNS_DECL_RE.findall(root_tag)}


def _ensure_root_namespace(root_tag: str, prefix: str, uri: str) -> str:
    if f"xmlns:{prefix}=" in root_tag:
        return root_tag
    return root_tag.replace(">", f' xmlns:{prefix}="{uri}">', 1)


def _normalize_xml_part_against_original(
    cleaned_xml: bytes,
    original_xml: bytes,
) -> bytes:
    cleaned_text = cleaned_xml.decode("utf-8")
    original_text = original_xml.decode("utf-8")
    cleaned_root_match = re.search(r"<[^?][^>]*>", cleaned_text)
    if cleaned_root_match is None:
        raise ValueError("No root tag found.")
    cleaned_root = cleaned_root_match.group(0)
    original_root = _first_root_tag(original_text)

    desired_prefix_by_uri = _namespace_prefixes(original_root)
    desired_prefix_by_uri.update(DOCUMENT_PREFIX_BY_URI)

    normalized_root = original_root

    body = cleaned_text[cleaned_root_match.end() :]
    cleaned_prefixes = _namespace_prefixes(cleaned_root)
    for uri, current_prefix in cleaned_prefixes.items():
        desired_prefix = desired_prefix_by_uri.get(uri)
        if not desired_prefix or desired_prefix == current_prefix:
            continue
        body = re.sub(
            rf"(</?){re.escape(current_prefix)}:",
            lambda match: f"{match.group(1)}{desired_prefix}:",
            body,
        )
        body = re.sub(
            rf"(\s){re.escape(current_prefix)}:",
            lambda match: f"{match.group(1)}{desired_prefix}:",
            body,
        )

    for uri, prefix in DOCUMENT_PREFIX_BY_URI.items():
        if f"{prefix}:" in body:
            normalized_root = _ensure_root_namespace(normalized_root, prefix, uri)
    for uri, prefix in cleaned_prefixes.items():
        if f"{prefix}:" in body:
            normalized_root = _ensure_root_namespace(normalized_root, prefix, uri)
    if "ns6:" in body:
        normalized_root = _ensure_root_namespace(
            normalized_root,
            "ns6",
            OFFICE_DRAWING_NAMESPACE,
        )

    return (normalized_root + body).encode("utf-8")


def _strip_run_shading_xml(xml_bytes: bytes) -> bytes:
    root = ET.fromstring(xml_bytes)
    for shd in list(root.findall(f".//{{{WORD_NAMESPACE}}}shd")):
        _remove_element_from_tree(root, shd)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _remove_element_from_tree(root, element) -> None:
    parent = _find_parent(root, element)
    if parent is not None:
        parent.remove(element)


def _run_has_only_breaks(run) -> bool:
    has_break = False
    for child in list(run):
        if child.tag == f"{{{WORD_NAMESPACE}}}br" or child.tag == f"{{{WORD_NAMESPACE}}}cr":
            has_break = True
            continue
        if child.tag == f"{{{WORD_NAMESPACE}}}rPr":
            continue
        return False
    return has_break


def _strip_leading_break_runs(root) -> None:
    for paragraph in root.findall(f".//{{{WORD_NAMESPACE}}}body/{{{WORD_NAMESPACE}}}p"):
        first_index = 0
        while first_index < len(paragraph) and paragraph[first_index].tag == f"{{{WORD_NAMESPACE}}}pPr":
            first_index += 1
        while (
            first_index < len(paragraph)
            and paragraph[first_index].tag == f"{{{WORD_NAMESPACE}}}r"
            and _run_has_only_breaks(paragraph[first_index])
        ):
            paragraph.remove(paragraph[first_index])


def _write_cleaned_package(
    *,
    input_path: Path,
    temp_output_path: Path,
    output_path: Path,
) -> None:
    final_temp_path = output_path.with_suffix(output_path.suffix + ".finaltmp")
    with ZipFile(input_path, "r") as zin, ZipFile(temp_output_path, "r") as ztemp:
        temp_names = set(ztemp.namelist())
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with ZipFile(final_temp_path, "w", compression=ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if (
                    (
                        _should_accept_revisions_in_part(info.filename)
                        or _should_clean_settings_part(info.filename)
                    )
                    and info.filename in temp_names
                ):
                    data = ztemp.read(info.filename)
                    if _should_accept_revisions_in_part(info.filename):
                        data = _strip_run_shading_xml(data)
                    if info.filename == "word/document.xml":
                        data = _normalize_xml_part_against_original(
                            data,
                            zin.read(info.filename),
                        )
                zout.writestr(info, data)
    final_temp_path.replace(output_path)


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _strip_paragraph_marking_formatting(paragraph) -> None:
    for element in list(paragraph._p.findall(".//w:highlight", {"w": WORD_NAMESPACE})):
        _remove_element(element)
    for element in list(paragraph._p.findall(".//w:shd", {"w": WORD_NAMESPACE})):
        _remove_element(element)


def _left_indent_pt(paragraph) -> float | None:
    indent = paragraph.paragraph_format.left_indent
    if indent is not None:
        return indent.pt
    style = paragraph.style
    if style is None or style.paragraph_format is None:
        return None
    style_indent = style.paragraph_format.left_indent
    if style_indent is None:
        return None
    return style_indent.pt


def _is_indented_content_paragraph(paragraph) -> bool:
    if not paragraph.text.strip():
        return False
    left_indent = _left_indent_pt(paragraph)
    return left_indent is not None and left_indent >= INDENT_THRESHOLD_PT


def _is_timestamp_paragraph(paragraph) -> bool:
    return bool(TIMESTAMP_LINE_RE.match(paragraph.text.strip()))


def _paragraph_has_drawing(paragraph) -> bool:
    return paragraph._p.find(
        ".//w:drawing",
        {"w": WORD_NAMESPACE},
    ) is not None


def _paragraph_has_content(paragraph) -> bool:
    return bool(paragraph.text.strip()) or _paragraph_has_drawing(paragraph)


def _section_kind(paragraph) -> str | None:
    text = paragraph.text.strip()
    if text in SUBTITLE_LABELS:
        return "subs"
    if text in SECTION_LABELS:
        return "other"
    return None


def _should_restore_yellow_highlight(paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return False
    if text in YELLOW_MARKER_TEXTS:
        return any(run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in paragraph.runs)
    if not (text.startswith("(") or text.endswith(")")):
        return False
    return any(run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in paragraph.runs)


def _normalize_highlights(doc: Document) -> None:
    had_yellow = [
        any(run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in paragraph.runs)
        for paragraph in doc.paragraphs
    ]
    paragraphs_to_restore: set[int] = set()
    for idx, paragraph in enumerate(doc.paragraphs):
        if _should_restore_yellow_highlight(paragraph):
            paragraphs_to_restore.add(idx)
            if idx > 0 and had_yellow[idx - 1]:
                paragraphs_to_restore.add(idx - 1)
        for run in paragraph.runs:
            rpr = run._r.rPr
            if rpr is not None:
                for shd in list(rpr.findall(qn("w:shd"))):
                    rpr.remove(shd)
            if (
                run.font.highlight_color is not None
                and run.font.highlight_color != WD_COLOR_INDEX.YELLOW
            ):
                run.font.highlight_color = None

    for idx in sorted(paragraphs_to_restore):
        paragraph = doc.paragraphs[idx]
        for run in paragraph.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _rewrite_revision_children(element) -> bool:
    changed = False
    idx = 0
    while idx < len(element):
        child = element[idx]
        if child.tag in INSERTION_TAGS:
            grandchildren = list(child)
            element.remove(child)
            for offset, grandchild in enumerate(grandchildren):
                element.insert(idx + offset, grandchild)
            changed = True
            idx += len(grandchildren)
            continue
        if child.tag in DELETION_TAGS or child.tag in CHANGE_TAGS:
            element.remove(child)
            changed = True
            continue
        if _rewrite_revision_children(child):
            changed = True
        idx += 1
    return changed


def _accept_revisions_in_tree(root) -> None:
    while _rewrite_revision_children(root):
        pass


def _disable_track_revisions_xml(xml_bytes: bytes) -> bytes:
    text = xml_bytes.decode("utf-8")
    return TRACK_REVISIONS_XML_RE.sub("", text).encode("utf-8")


def _accepted_revisions_docx_bytes(input_path: Path) -> bytes:
    buffer = BytesIO()
    with ZipFile(input_path, "r") as zin, ZipFile(buffer, "w", compression=ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if _should_accept_revisions_in_part(info.filename):
                data = _repair_word_xml_namespaces(data)
                root = ET.fromstring(data)
                _accept_revisions_in_tree(root)
                _strip_leading_break_runs(root)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif _should_clean_settings_part(info.filename):
                data = _disable_track_revisions_xml(data)
            zout.writestr(info, data)
    return buffer.getvalue()


def remove_sources_from_docx(input_path: Path, output_path: Path) -> None:
    doc = Document(BytesIO(_accepted_revisions_docx_bytes(input_path)))
    _normalize_highlights(doc)
    paragraphs = list(doc.paragraphs)
    remove_indexes = {
        idx
        for idx, paragraph in enumerate(paragraphs)
        if _is_indented_content_paragraph(paragraph)
    }
    if remove_indexes:
        block_starts = sorted(remove_indexes)
        block_ranges: list[tuple[int, int]] = []
        start = end = block_starts[0]
        for idx in block_starts[1:]:
            if idx == end + 1:
                end = idx
                continue
            block_ranges.append((start, end))
            start = end = idx
        block_ranges.append((start, end))

        for start, end in block_ranges:
            prev_idx = start - 1
            while prev_idx >= 0 and not _paragraph_has_content(paragraphs[prev_idx]):
                remove_indexes.add(prev_idx)
                prev_idx -= 1

            next_idx = end + 1
            while next_idx < len(paragraphs) and not _paragraph_has_content(paragraphs[next_idx]):
                remove_indexes.add(next_idx)
                next_idx += 1

    for index in sorted(remove_indexes, reverse=True):
        _remove_paragraph(paragraphs[index])

    paragraphs = list(doc.paragraphs)
    blank_indexes = []
    for idx, paragraph in enumerate(paragraphs):
        if _paragraph_has_content(paragraph):
            continue
        prev_non_blank = None
        next_non_blank = None
        for prev_idx in range(idx - 1, -1, -1):
            if _paragraph_has_content(paragraphs[prev_idx]):
                prev_non_blank = paragraphs[prev_idx]
                break
        for next_idx in range(idx + 1, len(paragraphs)):
            if _paragraph_has_content(paragraphs[next_idx]):
                next_non_blank = paragraphs[next_idx]
                break
        if prev_non_blank is None or next_non_blank is None:
            continue
        if _left_indent_pt(paragraph) is None:
            continue
        blank_indexes.append(idx)

    for index in reversed(blank_indexes):
        _remove_paragraph(paragraphs[index])

    paragraphs = list(doc.paragraphs)
    timestamp_blank_indexes = []
    for idx, paragraph in enumerate(paragraphs):
        if _paragraph_has_content(paragraph):
            continue
        if idx == 0 or idx + 1 >= len(paragraphs):
            continue
        if _is_timestamp_paragraph(paragraphs[idx - 1]) and _is_timestamp_paragraph(
            paragraphs[idx + 1]
        ):
            timestamp_blank_indexes.append(idx)

    for index in reversed(timestamp_blank_indexes):
        _remove_paragraph(paragraphs[index])

    paragraphs = list(doc.paragraphs)
    subtitle_blank_indexes = []
    current_section: str | None = None
    for idx, paragraph in enumerate(paragraphs):
        section = _section_kind(paragraph)
        if section is not None:
            current_section = section
            continue
        if current_section != "subs":
            continue
        if _paragraph_has_content(paragraph):
            continue
        prev_non_blank = None
        next_non_blank = None
        for prev_idx in range(idx - 1, -1, -1):
            if _paragraph_has_content(paragraphs[prev_idx]):
                prev_non_blank = paragraphs[prev_idx]
                break
        for next_idx in range(idx + 1, len(paragraphs)):
            next_section = _section_kind(paragraphs[next_idx])
            if next_section is not None:
                break
            if _paragraph_has_content(paragraphs[next_idx]):
                next_non_blank = paragraphs[next_idx]
                break
        if prev_non_blank is None or next_non_blank is None:
            continue
        if prev_non_blank.text.strip() in SUBTITLE_LABELS:
            continue
        subtitle_blank_indexes.append(idx)

    for index in reversed(subtitle_blank_indexes):
        _remove_paragraph(paragraphs[index])

    paragraphs = list(doc.paragraphs)
    repeated_blank_indexes = []
    previous_was_blank = False
    for idx, paragraph in enumerate(paragraphs):
        is_blank = not _paragraph_has_content(paragraph)
        if is_blank and previous_was_blank:
            repeated_blank_indexes.append(idx)
        previous_was_blank = is_blank

    for index in reversed(repeated_blank_indexes):
        _remove_paragraph(paragraphs[index])

    temp_output_path = output_path.with_suffix(output_path.suffix + ".tmp")
    doc.save(str(temp_output_path))
    _write_cleaned_package(
        input_path=input_path,
        temp_output_path=temp_output_path,
        output_path=output_path,
    )
    temp_output_path.unlink()


def _resolve_output_path(input_path: Path, output_path: str | None) -> Path:
    if output_path:
        return Path(output_path)
    return input_path


def _is_current_directory_docx(path: Path) -> bool:
    return path.is_file() and path.suffix.lower() == ".docx" and not path.name.startswith("~$")


def _resolve_input_paths(docx_files: list[str]) -> list[Path]:
    if docx_files:
        return [Path(raw_path) for raw_path in docx_files]
    return sorted(
        (
            path
            for path in Path.cwd().iterdir()
            if _is_current_directory_docx(path)
        ),
        key=lambda path: path.name.lower(),
    )


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Write cleaned DOCX files."
    )
    parser.add_argument(
        "docx_files",
        nargs="*",
        help="Input DOCX file paths. Default: all DOCX files in the current directory.",
    )
    parser.add_argument(
        "--output",
        help="Write cleaned content to a separate DOCX path. Default: overwrite input.",
    )
    args = parser.parse_args()

    if args.output and len(args.docx_files) != 1:
        parser.error("--output requires exactly one explicit input DOCX file.")

    input_paths = _resolve_input_paths(args.docx_files)
    if not input_paths:
        parser.error("no DOCX files found in the current directory.")

    for input_path in input_paths:
        output_path = _resolve_output_path(input_path, args.output)
        remove_sources_from_docx(input_path, output_path)
        print(f"[cleaned] {output_path}")


if __name__ == "__main__":
    main()
